"""Generate the transaction-workflow request document as .docx and .pdf.

Reads docs/blueprint/transaction_workflow_request.md as the source of truth and
produces matching Word and PDF files in the same folder.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
    Preformatted,
)
from reportlab.lib.enums import TA_LEFT


ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "docs" / "blueprint"
DOCX_PATH = OUT_DIR / "transaction_workflow_request.docx"
PDF_PATH = OUT_DIR / "transaction_workflow_request.pdf"


# -----------------------------------------------------------------------------
# Word
# -----------------------------------------------------------------------------

def _set_cell_shading(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def build_docx() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Request: End-to-End Transaction Workflow Specification")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    meta = [
        ("To", "Thomas"),
        ("From", "Monirul"),
        ("Date", "2026-05-05"),
        ("Re", "Defining the user-facing data entry workflow for each transaction type"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        p.add_run(value)

    doc.add_paragraph()

    # 1
    doc.add_heading("1. Why I need this", level=1)
    doc.add_paragraph(
        "The existing blueprint defines the rules a transaction must satisfy and "
        "the layers it passes through. To build the entry screens, AI proposal "
        "logic, and approval flows, I also need the step-by-step user journey per "
        "transaction type — what the preparer sees, what AI is allowed to "
        "propose, what is mandatory, who approves at what threshold, what happens "
        "on edge cases."
    )
    doc.add_paragraph(
        "Without that, I will guess — and a guess on workflow becomes a guess on "
        "accounting control, which I do not want to ship under your name."
    )

    # 2 — table
    doc.add_heading("2. What is already documented (so I don't ask twice)", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Source"
    hdr[1].text = "What it covers"
    hdr[2].text = "What it does not cover"
    for c in hdr:
        for para in c.paragraphs:
            for r in para.runs:
                r.bold = True
        _set_cell_shading(c, "1F3A5F")
        for para in c.paragraphs:
            for r in para.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    rows = [
        (
            "2026 04 17-DRAFT-CoA-Wealth management v2.xlsx",
            "3-layer model (Accounts / Masters / Posting rules); 09_Dimension_Validation_Rules gate; 3 canonical posting examples (school fees, equity buy, mortgage interest); per-account dimension requirements; controlled lists.",
            "The user journey to create those postings; approval routing; AI proposal scope; reversal / period rules; subledger workflow.",
        ),
        (
            "2026 04 30-Beakon-Architecture.pdf",
            "16-layer model; AI-assists / humans-approve / engine-validates separation; high-level Upload → Classify → Extract → Suggest flow.",
            "Per-transaction-type flow; thresholds; mandatory fields beyond the CoA gate; UX.",
        ),
        (
            "2026 04 17-Beakon Founder Working Paper.pdf",
            "Phase-1 build priorities, decision rules, \"must not be built yet\" list.",
            "Workflow at all.",
        ),
        (
            "thomas.ogg / thomas2.ogg / thomas3.ogg / thomas 4.ogg",
            "Voice memos on prior topics.",
            "If any of the questions below are already answered in these, please point me to the file — I will re-read before re-asking.",
        ),
    ]
    for src, covers, missing in rows:
        row = table.add_row().cells
        row[0].text = src
        row[1].text = covers
        row[2].text = missing

    # 3
    doc.add_heading("3. What I still need from you", level=1)
    doc.add_paragraph("A filled-in answer per transaction type, using the template in §5.")

    doc.add_heading("3.1 Transaction types — please confirm which are in Phase 1 scope", level=2)
    doc.add_paragraph("Tick / cross / \"later\":")
    types = [
        "General journal entry (manual)",
        "AP — supplier bill / receipt",
        "AR — client invoice / disbursement",
        "Bank / credit-card transaction + reconciliation",
        "Portfolio trade (buy / sell / corporate action)",
        "Loan transaction (drawdown, interest, repayment)",
        "Fixed asset (acquisition, depreciation, disposal)",
        "Payroll posting",
        "Period-end adjustments (accruals, prepayments, FX revaluation, IC elimination)",
        "Period close & lock",
    ]
    for i, t in enumerate(types, 1):
        doc.add_paragraph(f"☐  {i}. {t}", style="List Paragraph")
    doc.add_paragraph("☐  11. Anything I am missing? __________________________", style="List Paragraph")

    doc.add_paragraph(
        "If a type is \"later\", I will not build the entry screen yet — but I will "
        "still need to know whether the data model must accommodate it from day one."
    )

    # 4
    doc.add_heading("4. Cross-cutting questions (answer once, applies to all flows)", level=1)
    cc = [
        "Default approval threshold — at what amount / account class does a second approver kick in? Same threshold across entities, or entity-specific?",
        "Roles in scope for Phase 1 — which of {client user, accounting team, controller, senior reviewer, auditor (read-only), family-office manager, admin} can prepare? approve? override? I will use this to wire RBAC.",
        "AI citation policy — when AI proposes a treatment, must it cite (a) a prior similar posting, (b) the accounting standard + principle (per the per-entity standard rule), (c) the engine rule it is following, or all three?",
        "AI never-do list — what must AI never auto-propose under any flow? (e.g. prior-period postings, cross-entity / IC entries, anything touching tax lots, anything in a closed period.)",
        "Period rules — once a period is locked, are corrections done via (a) reopen + edit, (b) prior-period adjustment in current period, or (c) both, depending on materiality?",
        "FX rate source — single source per entity (e.g. ECB end-of-day), or per-transaction-type? Tolerance for manual override?",
        "Source-document mandatory? — is a linked document required to post, or only required to approve? Different per transaction type?",
        "Reversal pattern — full reversing entry vs. edit-with-audit-trail vs. void? Single rule or per-type?",
    ]
    for i, q in enumerate(cc, 1):
        doc.add_paragraph(f"{i}. {q}", style="List Number")

    # 5
    doc.add_heading("5. Template — please fill once per transaction type", level=1)
    doc.add_paragraph(
        "Copy-paste this block once per type and fill in. Short answers fine.",
        style="Intense Quote",
    )

    template_text = """Transaction type: ______________________________

A. Trigger / source document
   - What event starts the flow? (upload, bank feed, manual, scheduled, …)
   - Are source documents mandatory before posting? Before approving?
   - Where do the documents come from? (email, portal upload, integration)

B. Roles
   - Initiator role(s):
   - First-level approver:
   - Second-level approver (and the threshold that triggers it):
   - Override / exception authority:

C. Mandatory fields beyond the CoA workbook rules
   - The workbook already enforces per-account dimensions. List anything additional this transaction type needs, e.g.:
       AP    → supplier ID, due date, payment method, VAT code
       Trade → trade date, settlement date, custodian, tax-lot method (FIFO/LIFO/avg)
   - Field-level defaults the system should pre-fill:

D. AI proposal scope
   - What MAY AI propose automatically?
   - What MUST AI never propose for this type?
   - What MUST every AI proposal cite? (prior entry / accounting standard / engine rule)

E. Engine validation (beyond debits=credits + dimension completeness)
   - Examples: 3-way match for AP, tax-lot availability for sell, open-period check,
     FX-rate-within-tolerance, loan-balance-not-negative.

F. Posting outcomes
   - Subledgers touched: AP / AR / Bank / FA / Tax Lots / Loan / Inventory / …
   - Reversal / cancel path:
   - Source-document linking — automatic or manual?

G. Edge cases that need explicit handling
   - Multi-currency / FX rate override
   - Multi-entity / intercompany side-effects
   - Rebillable flag → what additional posting fires?
   - Partial payments / partial matches
   - Missing-document state — queue, alert, or block?
   - Anything else specific to this type:

H. UX preferences (optional — I will choose if you have no preference)
   - Single-screen vs. wizard
   - AI proposal as pre-filled fields vs. side-by-side proposal/final view
   - Save-draft vs. submit-for-approval — separate states?
"""
    p = doc.add_paragraph()
    run = p.add_run(template_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)

    # 6
    doc.add_heading("6. One worked example I would like first", level=1)
    doc.add_paragraph(
        "If you only have time for one, the AP supplier invoice flow — the same "
        "example used in the architecture PDF (\"From One Supplier Invoice to "
        "Reviewed Financial Reporting\"). If you walk me through that one "
        "end-to-end with concrete numbers, dimensions, approval thresholds, and "
        "edge cases, I can extrapolate the others and bring them back to you for "
        "sign-off rather than starting from a blank page each time."
    )
    doc.add_paragraph("A 30-minute live walkthrough or a single voice memo would be ideal.")

    # 7
    doc.add_heading("7. How to deliver — whatever is fastest for you", level=1)
    for line in [
        "Annotate this document directly (Word / PDF / handwritten — any of those work).",
        "Voice memo per transaction type (same format as thomas.ogg); I will transcribe and structure.",
        "One live walkthrough on AP, then I draft the rest and you review.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # 8
    doc.add_heading("8. What I will produce from your answer", level=1)
    for line in [
        "docs/blueprint/transaction_workflows.md — companion to instructions_dump.md, one section per transaction type, in the same verbatim-blueprint style.",
        "A small per-flow diagram (preparer → AI proposal → engine validation → approver → post → reporting) for each type, for sign-off before I implement.",
        "A mapping from each flow to the 16 layers, so we can confirm no layer boundary is being crossed.",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_paragraph(
        "I will not start UI implementation on any transaction type until the workflow for it is signed off."
    )

    # Final ask
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("One ask: ")
    r.bold = True
    p.add_run(
        "if the answer to any question is \"I have not decided yet\", please mark "
        "it that way explicitly. That is more useful to me than a tentative answer "
        "I might lock into the schema and have to undo later."
    )

    doc.save(DOCX_PATH)


# -----------------------------------------------------------------------------
# PDF
# -----------------------------------------------------------------------------

def build_pdf() -> None:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleX",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=18,
        textColor=colors.HexColor("#1F3A5F"),
        alignment=TA_LEFT,
        spaceAfter=14,
    )
    h1 = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=14,
        textColor=colors.HexColor("#1F3A5F"),
        spaceBefore=14,
        spaceAfter=8,
    )
    h2 = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        textColor=colors.HexColor("#1F3A5F"),
        spaceBefore=10,
        spaceAfter=6,
    )
    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        spaceAfter=6,
    )
    bullet = ParagraphStyle(
        "Bullet",
        parent=body,
        leftIndent=14,
        bulletIndent=2,
    )
    mono = ParagraphStyle(
        "Mono",
        parent=body,
        fontName="Courier",
        fontSize=8.5,
        leading=11,
    )

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=2.0 * cm,
        rightMargin=2.0 * cm,
        topMargin=2.0 * cm,
        bottomMargin=2.0 * cm,
        title="Transaction Workflow Request",
        author="Monirul",
    )
    story = []
    story.append(Paragraph("Request: End-to-End Transaction Workflow Specification", title_style))
    for label, value in [
        ("To", "Thomas"),
        ("From", "Monirul"),
        ("Date", "2026-05-05"),
        ("Re", "Defining the user-facing data entry workflow for each transaction type"),
    ]:
        story.append(Paragraph(f"<b>{label}:</b> {value}", body))
    story.append(Spacer(1, 0.3 * cm))

    story.append(Paragraph("1. Why I need this", h1))
    story.append(Paragraph(
        "The existing blueprint defines the <b>rules</b> a transaction must satisfy "
        "and the <b>layers</b> it passes through. To build the entry screens, AI "
        "proposal logic, and approval flows, I also need the <b>step-by-step user "
        "journey</b> per transaction type — what the preparer sees, what AI is "
        "allowed to propose, what is mandatory, who approves at what threshold, "
        "what happens on edge cases.",
        body,
    ))
    story.append(Paragraph(
        "Without that, I will guess — and a guess on workflow becomes a guess on "
        "accounting control, which I do not want to ship under your name.",
        body,
    ))

    story.append(Paragraph("2. What is already documented (so I don't ask twice)", h1))
    table_data = [
        ["Source", "What it covers", "What it does not cover"],
        [
            "CoA workbook (2026 04 17 v2.xlsx)",
            "3-layer model; dimension validation gate; 3 canonical posting examples; per-account dimension requirements; controlled lists.",
            "The user journey to create postings; approval routing; AI proposal scope; reversal / period rules; subledger workflow.",
        ],
        [
            "Architecture PDF (2026 04 30)",
            "16-layer model; AI-assists / humans-approve / engine-validates separation; high-level Upload → Classify → Extract → Suggest flow.",
            "Per-transaction-type flow; thresholds; mandatory fields beyond the CoA gate; UX.",
        ],
        [
            "Founder Working Paper (2026 04 17)",
            "Phase-1 build priorities, decision rules, \"must not be built yet\" list.",
            "Workflow at all.",
        ],
        [
            "thomas.ogg / 2 / 3 / 4",
            "Voice memos on prior topics.",
            "If any question below is already answered in these, point me to the file — I will re-read before re-asking.",
        ],
    ]
    wrapped = [[Paragraph(cell, body) for cell in row] for row in table_data]
    tbl = Table(wrapped, colWidths=[4.5 * cm, 6.0 * cm, 6.0 * cm], repeatRows=1)
    tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F3A5F")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#999999")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(tbl)

    story.append(Paragraph("3. What I still need from you", h1))
    story.append(Paragraph("A filled-in answer per transaction type, using the template in §5.", body))

    story.append(Paragraph("3.1 Transaction types — please confirm which are in Phase 1 scope", h2))
    story.append(Paragraph("Tick / cross / \"later\":", body))
    types = [
        "General journal entry (manual)",
        "AP — supplier bill / receipt",
        "AR — client invoice / disbursement",
        "Bank / credit-card transaction + reconciliation",
        "Portfolio trade (buy / sell / corporate action)",
        "Loan transaction (drawdown, interest, repayment)",
        "Fixed asset (acquisition, depreciation, disposal)",
        "Payroll posting",
        "Period-end adjustments (accruals, prepayments, FX revaluation, IC elimination)",
        "Period close &amp; lock",
    ]
    for i, t in enumerate(types, 1):
        story.append(Paragraph(f"☐ &nbsp; {i}. {t}", bullet))
    story.append(Paragraph("☐ &nbsp; 11. Anything I am missing? __________________________", bullet))

    story.append(Paragraph(
        "If a type is \"later\", I will not build the entry screen yet — but I will "
        "still need to know whether the data model must accommodate it from day one.",
        body,
    ))

    story.append(Paragraph("4. Cross-cutting questions (answer once, applies to all flows)", h1))
    cc = [
        "Default approval threshold — at what amount / account class does a second approver kick in? Same threshold across entities, or entity-specific?",
        "Roles in scope for Phase 1 — which of {client user, accounting team, controller, senior reviewer, auditor (read-only), family-office manager, admin} can prepare? approve? override?",
        "AI citation policy — when AI proposes a treatment, must it cite (a) a prior similar posting, (b) the accounting standard + principle, (c) the engine rule it is following, or all three?",
        "AI never-do list — what must AI never auto-propose under any flow? (prior-period postings, cross-entity / IC entries, anything touching tax lots, anything in a closed period, etc.)",
        "Period rules — once a period is locked, are corrections done via (a) reopen + edit, (b) prior-period adjustment in current period, or (c) both, depending on materiality?",
        "FX rate source — single source per entity (e.g. ECB end-of-day), or per-transaction-type? Tolerance for manual override?",
        "Source-document mandatory? — required to post, or only to approve? Different per transaction type?",
        "Reversal pattern — full reversing entry vs. edit-with-audit-trail vs. void? Single rule or per-type?",
    ]
    for i, q in enumerate(cc, 1):
        story.append(Paragraph(f"{i}. {q}", body))

    story.append(PageBreak())

    story.append(Paragraph("5. Template — please fill once per transaction type", h1))
    story.append(Paragraph(
        "<i>Copy-paste this block once per type and fill in. Short answers fine.</i>",
        body,
    ))
    template_text = """Transaction type: ______________________________

A. Trigger / source document
   - What event starts the flow? (upload, bank feed, manual, scheduled, ...)
   - Are source documents mandatory before posting? Before approving?
   - Where do the documents come from? (email, portal upload, integration)

B. Roles
   - Initiator role(s):
   - First-level approver:
   - Second-level approver (and the threshold that triggers it):
   - Override / exception authority:

C. Mandatory fields beyond the CoA workbook rules
   - The workbook already enforces per-account dimensions. List anything
     additional this transaction type needs, e.g.:
       AP    -> supplier ID, due date, payment method, VAT code
       Trade -> trade date, settlement date, custodian, tax-lot method
   - Field-level defaults the system should pre-fill:

D. AI proposal scope
   - What MAY AI propose automatically?
   - What MUST AI never propose for this type?
   - What MUST every AI proposal cite? (prior entry / standard / engine rule)

E. Engine validation (beyond debits=credits + dimension completeness)
   - Examples: 3-way match for AP, tax-lot availability for sell,
     open-period check, FX-rate-within-tolerance, loan-balance-not-negative.

F. Posting outcomes
   - Subledgers touched: AP / AR / Bank / FA / Tax Lots / Loan / Inventory / ...
   - Reversal / cancel path:
   - Source-document linking - automatic or manual?

G. Edge cases that need explicit handling
   - Multi-currency / FX rate override
   - Multi-entity / intercompany side-effects
   - Rebillable flag -> what additional posting fires?
   - Partial payments / partial matches
   - Missing-document state - queue, alert, or block?
   - Anything else specific to this type:

H. UX preferences (optional - I will choose if you have no preference)
   - Single-screen vs. wizard
   - AI proposal as pre-filled fields vs. side-by-side proposal/final view
   - Save-draft vs. submit-for-approval - separate states?
"""
    story.append(Preformatted(template_text, mono))

    story.append(Paragraph("6. One worked example I would like first", h1))
    story.append(Paragraph(
        "If you only have time for one, the <b>AP supplier invoice</b> flow — the "
        "same example used in the architecture PDF (\"From One Supplier Invoice to "
        "Reviewed Financial Reporting\"). If you walk me through that one "
        "end-to-end with concrete numbers, dimensions, approval thresholds, and "
        "edge cases, I can extrapolate the others and bring them back to you for "
        "sign-off rather than starting from a blank page each time.",
        body,
    ))
    story.append(Paragraph(
        "A 30-minute live walkthrough or a single voice memo would be ideal.",
        body,
    ))

    story.append(Paragraph("7. How to deliver — whatever is fastest for you", h1))
    for line in [
        "Annotate this document directly (Word / PDF / handwritten — any of those work).",
        "Voice memo per transaction type (same format as thomas.ogg); I will transcribe and structure.",
        "One live walkthrough on AP, then I draft the rest and you review.",
    ]:
        story.append(Paragraph("• " + line, bullet))

    story.append(Paragraph("8. What I will produce from your answer", h1))
    for line in [
        "<b>docs/blueprint/transaction_workflows.md</b> — companion to instructions_dump.md, one section per transaction type, in the same verbatim-blueprint style.",
        "A small per-flow diagram (preparer → AI proposal → engine validation → approver → post → reporting) for each type, for sign-off before I implement.",
        "A mapping from each flow to the 16 layers, so we can confirm no layer boundary is being crossed.",
    ]:
        story.append(Paragraph("• " + line, bullet))
    story.append(Paragraph(
        "I will not start UI implementation on any transaction type until the workflow for it is signed off.",
        body,
    ))

    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph(
        "<b>One ask:</b> if the answer to any question is \"I have not decided yet\", "
        "please mark it that way explicitly. That is more useful to me than a "
        "tentative answer I might lock into the schema and have to undo later.",
        body,
    ))

    doc.build(story)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_pdf()
    print(f"Wrote: {DOCX_PATH}")
    print(f"Wrote: {PDF_PATH}")


if __name__ == "__main__":
    main()
