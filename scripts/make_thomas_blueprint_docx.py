"""Generate the Beakon Engine Demonstration document for Thomas.

Output: test_data/beakon_engine_demonstration.docx — a Word document that
Thomas can read on any device, write comments in, and return as the
canonical blueprint for the engine going forward.

This is NOT a developer demo script. Operational details (port numbers,
shell commands, env vars, "if Thomas asks…" callouts) are stripped.
What remains is a customer-facing description of what Beakon does, why
it does it that way, and what each demonstration proves.

Run:
    venv\\Scripts\\python.exe scripts\\make_thomas_blueprint_docx.py
"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent.parent / "test_data"
OUT_DIR.mkdir(exist_ok=True)
DOCX_PATH = OUT_DIR / "beakon_engine_demonstration.docx"


# ── Brand palette ───────────────────────────────────────────────────────
BRAND = RGBColor(0x1A, 0x4F, 0x4A)
BRAND_LIGHT_HEX = "E8F0EE"
GRAY_900 = RGBColor(0x11, 0x18, 0x27)
GRAY_700 = RGBColor(0x37, 0x41, 0x51)
GRAY_500 = RGBColor(0x6B, 0x72, 0x80)
GRAY_LIGHT_HEX = "F3F4F6"


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "D1D5DB")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def style_run(run, *, bold=False, size=11, color=None, italic=False, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text, *, size=11, color=None, bold=False, italic=False,
                  align=None, space_after=4, indent=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.left_indent = indent
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    style_run(run, bold=bold, size=size, color=color or GRAY_700, italic=italic)
    return p


def add_heading(doc, text, *, level=1):
    sizes = {0: 28, 1: 18, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level <= 1 else 12)
    p.paragraph_format.space_after = Pt(8 if level <= 1 else 6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, bold=True, size=sizes.get(level, 11), color=BRAND)
    return p


def add_bullet(doc, text, *, indent=Cm(0.5), bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = indent
    if bold_lead:
        run = p.add_run(bold_lead)
        style_run(run, bold=True, size=11, color=GRAY_900)
        run2 = p.add_run(" " + text)
        style_run(run2, size=11, color=GRAY_700)
    else:
        run = p.add_run(text)
        style_run(run, size=11, color=GRAY_700)
    return p


def add_callout(doc, text, *, fill=BRAND_LIGHT_HEX):
    """Boxed paragraph for important notes — single-cell table for shading."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    set_cell_borders(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    style_run(run, size=10, color=GRAY_900)
    # Spacer after
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    return table


def add_demo_table(doc, rows):
    """A 'What you'll see / What it proves' two-column table."""
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.autofit = True

    # Header
    hdr = table.rows[0].cells
    for i, label in enumerate(["What you'll see", "What it proves"]):
        shade_cell(hdr[i], "1A4F4A")
        set_cell_borders(hdr[i])
        para = hdr[i].paragraphs[0]
        run = para.add_run(label)
        style_run(run, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))

    # Body
    for i, (col_a, col_b) in enumerate(rows, start=1):
        for j, content in enumerate((col_a, col_b)):
            cell = table.rows[i].cells[j]
            set_cell_borders(cell)
            if i % 2 == 0:
                shade_cell(cell, GRAY_LIGHT_HEX)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(content)
            style_run(run, size=10, color=GRAY_700)

    # Spacer after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    return table


def add_workflow_chart(doc, steps, *, current_index=None):
    """Render a horizontal workflow as a single-row table of shaded
    cells separated by arrow cells. ``current_index`` highlights one
    step in brand color (None = no highlight, all neutral)."""
    n = len(steps)
    # Layout: step / arrow / step / arrow / ... / step
    n_cols = n * 2 - 1
    table = doc.add_table(rows=1, cols=n_cols)
    table.autofit = True

    for i in range(n_cols):
        cell = table.rows[0].cells[i]
        is_step = (i % 2 == 0)
        step_idx = i // 2
        if is_step:
            highlighted = current_index is not None and step_idx == current_index
            shade_cell(cell, "1A4F4A" if highlighted else "E8F0EE")
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(steps[step_idx])
            style_run(
                run, bold=True, size=10,
                color=RGBColor(0xFF, 0xFF, 0xFF) if highlighted else BRAND,
            )
        else:
            # Arrow cell — narrow, no border, just a centred chevron.
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run("→")
            style_run(run, bold=True, size=14, color=GRAY_500)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    return table


def add_example_table(doc, headers, rows, *, footer=None):
    """Render a worked-example data table with brand-coloured header,
    alternating row shading, and an optional bold footer (for totals)."""
    total_rows = len(rows) + 1 + (1 if footer else 0)
    table = doc.add_table(rows=total_rows, cols=len(headers))
    table.autofit = True

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, "1A4F4A")
        set_cell_borders(cell)
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(h)
        style_run(run, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))

    # Body
    for r, row_vals in enumerate(rows, start=1):
        for c, val in enumerate(row_vals):
            cell = table.rows[r].cells[c]
            set_cell_borders(cell)
            if r % 2 == 0:
                shade_cell(cell, GRAY_LIGHT_HEX)
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(str(val))
            style_run(run, size=10, color=GRAY_700)

    # Footer (totals row, bold + brand-tinted)
    if footer:
        footer_row = total_rows - 1
        for c, val in enumerate(footer):
            cell = table.rows[footer_row].cells[c]
            shade_cell(cell, BRAND_LIGHT_HEX)
            set_cell_borders(cell)
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(str(val))
            style_run(run, bold=True, size=10, color=BRAND)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    return table


def add_faq(doc, q, a):
    """A single Q/A pair — bold question in brand color, gray answer."""
    qp = doc.add_paragraph()
    qp.paragraph_format.space_before = Pt(8)
    qp.paragraph_format.space_after = Pt(2)
    qp.paragraph_format.keep_with_next = True
    q_run = qp.add_run("Q. " + q)
    style_run(q_run, bold=True, size=11, color=BRAND)

    ap = doc.add_paragraph()
    ap.paragraph_format.space_after = Pt(2)
    a_run = ap.add_run("A. " + a)
    style_run(a_run, size=11, color=GRAY_700)


def add_review_box(doc, prompt):
    """Empty space for Thomas to write comments. Visible bordered box
    with a prompt line at the top."""
    table = doc.add_table(rows=2, cols=1)
    table.autofit = True
    cell_top = table.rows[0].cells[0]
    cell_body = table.rows[1].cells[0]
    shade_cell(cell_top, "F4ED E8")  # warm cream — invites markup
    set_cell_borders(cell_top)
    set_cell_borders(cell_body)

    p_top = cell_top.paragraphs[0]
    run = p_top.add_run("Thomas's notes — " + prompt)
    style_run(run, bold=True, size=9, italic=True, color=GRAY_500)

    # ~3 blank lines for handwritten or typed comments
    body_para = cell_body.paragraphs[0]
    body_para.paragraph_format.space_before = Pt(2)
    body_para.paragraph_format.space_after = Pt(2)
    for _ in range(3):
        cell_body.add_paragraph()

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


# ── Build ───────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page setup — narrow margins for tighter pages
    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Default font — body text Calibri 11pt
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Cover ───────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("Beakon")
    style_run(title_run, bold=True, size=36, color=BRAND, font="Calibri")
    title.paragraph_format.space_after = Pt(0)

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("Engine Demonstration & Blueprint Document")
    style_run(sub_run, size=16, color=GRAY_700, italic=True)
    subtitle.paragraph_format.space_after = Pt(24)

    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"Prepared for Thomas Allina  ·  "
        f"{date.today().strftime('%B %Y')}  ·  Draft for review"
    )
    style_run(meta_run, size=10, color=GRAY_500)
    meta.paragraph_format.space_after = Pt(28)

    add_callout(doc,
        "This document describes what Beakon's accounting engine does, "
        "why each capability exists, and what eight separate demonstrations "
        "prove about the engine's correctness. It is intended as a working "
        "blueprint — please mark up any section, write notes inline, and "
        "return it. Your annotations will become the canonical record of "
        "what Beakon should be.",
    )

    # ── Section 1: What Beakon is ──────────────────────────────────
    add_heading(doc, "1. What Beakon is", level=1)

    add_paragraph(doc,
        "Beakon is an AI-native finance operating system for family offices and "
        "their controlled entities. The engine is built around three principles "
        "drawn from your founder working paper:",
        size=11, color=GRAY_700, space_after=8,
    )
    add_bullet(doc,
        "every accounting record posts as a balanced double-entry journal "
        "entry — nothing else can land on the ledger;",
        bold_lead="Double-entry first.",
    )
    add_bullet(doc,
        "the AI may read documents and draft journal entries, but a human "
        "must approve before anything posts; segregation of duties is enforced "
        "at the engine level, not as a policy afterthought;",
        bold_lead="AI proposes, humans approve.",
    )
    add_bullet(doc,
        "every figure on every report can be drilled back to its journal "
        "entry, and that journal entry back to its source document plus the "
        "AI metadata that drafted it. There is no opaque step.",
        bold_lead="Every cent traceable.",
    )

    add_paragraph(doc,
        "The current build covers three intake flows (vendor bills, customer "
        "invoices, and bank statement feeds), the full approval state machine, "
        "six financial reports, and a unified bookkeeper review queue. Multi-"
        "currency and multi-entity are working today.",
        size=11, color=GRAY_700, space_after=12,
    )

    add_review_box(doc, "anything missing from this overview?")

    # ── Section 2: What's been built ────────────────────────────────
    add_heading(doc, "2. What's been built", level=1)

    add_heading(doc, "Three intake flows", level=2)
    add_paragraph(doc,
        "Every accounting transaction enters through one of three doors. "
        "All three converge into balanced journal entries that feed the same "
        "ledger.",
    )

    add_bullet(doc,
        "the bookkeeper uploads a vendor receipt as PDF or image. AI extracts "
        "vendor name, invoice number, dates, line items, totals, currency, and "
        "suggests an expense account from the entity's chart. The reviewer "
        "sees the AI's proposal in a draft, edits as needed, and approves. "
        "Approval auto-posts the accrual journal entry.",
        bold_lead="Vendor bills (AP).",
    )
    add_bullet(doc,
        "same upload flow, but AI suggests a revenue account and matches the "
        "customer record. Useful for importing legacy invoices issued before "
        "Beakon went live.",
        bold_lead="Customer invoices (AR).",
    )
    add_bullet(doc,
        "CSV statement import. Each transaction can be categorised in one "
        "click; AI suggests the offset account based on the description, "
        "amount, and the entity's chart of accounts. The resulting journal "
        "entry auto-submits for approval.",
        bold_lead="Bank statement feed.",
    )

    add_heading(doc, "The approval state machine", level=2)
    add_paragraph(doc,
        "Every journal entry — regardless of where it came from — flows "
        "through the same four-step lifecycle:",
    )
    add_paragraph(doc,
        "draft  →  pending approval  →  approved  →  posted",
        size=12, color=BRAND, bold=True, italic=True, space_after=8,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(doc,
        "Two different humans must touch every posted entry — the submitter "
        "cannot approve their own work. This is enforced at the engine, not "
        "at the role layer, so it cannot be turned off by reconfiguring "
        "permissions. Bills and invoices follow the same lifecycle, with the "
        "approval action automatically generating the underlying journal entry.",
        space_after=12,
    )

    add_heading(doc, "AI assistance, not autonomy", level=2)
    add_paragraph(doc,
        "AI features in Beakon are read-and-suggest, never read-and-act. The "
        "engine surfaces an AI's confidence score and reasoning on every "
        "draft, and records the model used in the audit trail. The reviewer "
        "is always the last hand on the record.",
        space_after=4,
    )
    add_paragraph(doc,
        "Today the AI features are: bill OCR + extraction, invoice OCR + "
        "extraction, bank-line categorisation, and short executive commentary "
        "on each report. All four can run on local infrastructure (privacy-"
        "first, nothing leaves the office) or on Anthropic's Claude API for "
        "higher accuracy on hard-to-read documents.",
        space_after=12,
    )

    add_review_box(doc, "what's missing or mis-prioritised?")

    # ── Section 2.5: How a transaction flows (workflow charts) ──────
    add_heading(doc, "3. How a transaction flows", level=1)
    add_paragraph(doc,
        "The same lifecycle governs every record in Beakon. The diagram "
        "below shows the four states a journal entry passes through, from "
        "creation to the ledger:",
        space_after=8,
    )
    add_workflow_chart(doc, [
        "Draft", "Pending approval", "Approved", "Posted",
    ])
    add_paragraph(doc,
        "Movement between states requires an explicit human action. There "
        "is no auto-posting, no scheduled batch, no override. A draft can "
        "also be cancelled at any time before approval — cancelled records "
        "stay in the system for audit but never reach the ledger.",
        space_after=14,
    )

    add_heading(doc, "Vendor bill (AP) workflow", level=2)
    add_paragraph(doc,
        "An incoming vendor invoice. The bookkeeper uploads the PDF or "
        "image; AI drafts the bill; a second human approves; the accrual "
        "journal entry posts; the bill becomes payable.",
        space_after=8,
    )
    add_workflow_chart(doc, [
        "Receipt uploaded",
        "AI extraction",
        "Bookkeeper review",
        "Approved",
        "JE posted",
    ])
    add_paragraph(doc,
        "Customer invoices (AR) and bank-feed transactions follow the same "
        "shape — the only differences are which document is uploaded, which "
        "side of the journal entry the AI suggests, and where the resulting "
        "posting lands on the chart of accounts.",
        space_after=14,
    )

    add_heading(doc, "End-to-end traceability", level=2)
    add_paragraph(doc,
        "Any number on any report can be traced back through the system to "
        "its origin. The chain looks like this:",
        space_after=8,
    )
    add_workflow_chart(doc, [
        "Trial balance figure",
        "Account ledger row",
        "Journal entry",
        "Source document",
        "AI metadata",
    ])
    add_paragraph(doc,
        "Each arrow is a single click in the user interface. There is no "
        "report without a corresponding ledger entry, no ledger entry "
        "without a source document, and no AI-drafted record without the "
        "model name, confidence score, and accounting-standard reasoning "
        "preserved on the audit trail.",
        space_after=12,
    )

    add_review_box(doc, "any step in these workflows you would change?")

    # ── Section: Eight demonstrations ────────────────────────────────
    add_heading(doc, "4. The eight demonstrations", level=1)
    add_paragraph(doc,
        "Each demonstration probes a different property of the engine. "
        "Together they prove the system is internally consistent — the "
        "same transactions produce the same numbers no matter which "
        "report you view them through.",
        space_after=10,
    )

    # 4.1
    add_heading(doc, "4.1 Trial Balance", level=2)
    add_paragraph(doc,
        "The most basic test in accounting: do total debits equal total "
        "credits? If yes, every journal entry in the system is a balanced "
        "double-entry. The engine exposes this as a green Balanced badge "
        "on the report header.",
    )
    add_demo_table(doc, [
        ("Trial Balance with green Balanced badge.",
         "Every JE posted to the ledger respects double-entry."),
        ("Same account code shown per entity (e.g. five Operating Bank rows).",
         "Multi-entity accounting is working — each entity carries its own "
         "instance of the chart, no silent merging."),
        ("Export to CSV downloads cleanly into Excel.",
         "Reports are file-shareable, not just on-screen artefacts."),
    ])

    # 4.2
    add_heading(doc, "4.2 Balance Sheet", level=2)
    add_paragraph(doc,
        "The accounting equation must hold at any reporting date: "
        "Assets = Liabilities + Equity. The badge turns red if the "
        "engine fails to maintain this — a fundamental sanity check.",
    )
    add_demo_table(doc, [
        ("Three sections with totals: Assets, Liabilities, Equity.",
         "Standard presentation matching IFRS and US-GAAP conventions."),
        ("Green Balanced badge with no rounding diff.",
         "The accounting equation holds at the as-of date."),
        ("Date selector + quick-period chips (Today / End of last month / Quarter / Year).",
         "Reports are point-in-time queryable, not just live snapshots."),
    ])

    # 4.3
    add_heading(doc, "4.3 Profit & Loss", level=2)
    add_paragraph(doc,
        "Operations flowing through. Bills become expenses, invoices become "
        "revenue, bank fees become expenses. The P&L should reflect every "
        "approved transaction in the period.",
    )
    add_demo_table(doc, [
        ("Sectioned layout: Revenue / COGS / Gross Profit / Operating Expenses / "
         "Operating Income / Other / Net Income.",
         "Standard multi-step income statement format."),
        ("Each row shows '% of revenue' alongside the figure.",
         "Common-size analysis for proportional comparison across periods."),
        ("Negative numbers in parentheses (accounting convention).",
         "Typography and presentation match what your auditor expects to see."),
    ])

    # 4.4
    add_heading(doc, "4.4 Cash Flow Statement", level=2)
    add_paragraph(doc,
        "The strictest test. Cash flow is reconstructed from journal entries "
        "and must reconcile to the balance sheet's cash position. The "
        "Reconciles to balance sheet badge is the proof that all four reports "
        "(TB, BS, P&L, CF) agree with each other.",
    )
    add_demo_table(doc, [
        ("Direct-method statement: Operating, Investing, Financing.",
         "Standard presentation; preferred over indirect for small to "
         "mid-sized entities by IFRS guidance."),
        ("Green Reconciles to balance sheet badge.",
         "The four core reports are mutually consistent — strongest single "
         "proof that the engine is internally correct."),
        ("Empty-state hint surfaces if no cash activity in the period.",
         "The system is honest about emptiness instead of showing zeroes "
         "across an unhelpful table."),
    ])

    # 4.5
    add_heading(doc, "4.5 AP and AR Aging", level=2)
    add_paragraph(doc,
        "Daily working-capital views. Who owes the entities money, who they "
        "owe, and how old each obligation is. Color-coded by age bucket so "
        "anything past due jumps off the page.",
    )
    add_demo_table(doc, [
        ("Vendor / customer rows with age bucket columns "
         "(Current / 1-30 / 31-60 / 61-90 / 90+).",
         "Standard aging report structure used by every accountant."),
        ("Click a row to drill into individual outstanding documents.",
         "Aging is queryable, not just a summary — supports collection / "
         "payment workflows."),
        ("Native amount displayed alongside reporting-currency amount.",
         "Multi-currency receivables / payables don't lose their original "
         "denomination."),
    ])

    # 4.6
    add_heading(doc, "4.6 End-to-end traceability", level=2)
    add_paragraph(doc,
        "The deepest proof. Starting from a single number on the trial balance, "
        "you can drill through the system to find the journal entry that "
        "produced it, the source document that produced the journal entry, "
        "and the AI metadata that produced the draft. Nothing is opaque.",
    )
    add_demo_table(doc, [
        ("Click an account on Trial Balance → open Account Ledger.",
         "Every report number leads to its underlying detail."),
        ("Click a JE on the Account Ledger → open JE detail page.",
         "Detail includes balanced lines, status timeline, AI reasoning, "
         "linked source document."),
        ("Click the linked source document.",
         "The original bill or bank line opens. Audit chain is complete."),
    ])

    # 4.7
    add_heading(doc, "4.7 The unified review queue", level=2)
    add_paragraph(doc,
        "One screen for everything pending sign-off — journal entries, "
        "bills, invoices. Color-coded by age so nothing falls through "
        "the cracks. Multi-entity controllers see only their slice.",
    )
    add_demo_table(doc, [
        ("Pending bucket lists items from all three intake flows.",
         "The bookkeeper's daily inbox is unified, not three separate queues."),
        ("Inline approve / reject buttons on each row.",
         "Routine approvals require one click."),
        ("Filter by entity for multi-entity controllers.",
         "Each reviewer sees only the entities they're permissioned on."),
    ])

    # 4.8
    add_heading(doc, "4.8 The audit log", level=2)
    add_paragraph(doc,
        "Append-only record of every action. Who did what, when, with what "
        "metadata. Regulator-ready: nothing is erasable, nothing happens "
        "without a corresponding event.",
    )
    add_demo_table(doc, [
        ("Stream of events with actor, object, action, timestamp.",
         "Standard audit trail meeting common-framework requirements."),
        ("AI events tagged with the model that produced the draft.",
         "AI is recorded as an actor type, not hidden behind the user "
         "who accepted the suggestion."),
        ("Approval events show submitter and approver are different users.",
         "Segregation-of-duties evidence preserved on every posted entry."),
    ])

    add_review_box(doc, "which demonstration shows the wrong thing?")

    # ── Section: Worked examples ──────────────────────────────────
    add_heading(doc, "5. Worked examples", level=1)
    add_paragraph(doc,
        "Three concrete scenarios. Each one shows the user-facing inputs, "
        "the journal entry the engine creates, and the audit trail it "
        "leaves behind. Identical to what your bookkeeper will see in "
        "production.",
        space_after=10,
    )

    # Example 1: Vendor bill upload
    add_heading(doc, "5.1 Vendor bill — office supplies (USD → CHF)", level=2)
    add_paragraph(doc,
        "BlueSpruce Office Supplies sends Thomas Foundation a bill for "
        "$258.39 USD. The bookkeeper drags the PDF into the New Bill "
        "drawer. Claude Haiku reads it in 3 seconds and prefills the form.",
        space_after=8,
    )
    add_paragraph(doc, "AI extraction:", bold=True, size=10, space_after=2)
    add_example_table(doc,
        ["Field", "Extracted value", "Source"],
        [
            ("Vendor name", "BlueSpruce Office Supplies, Inc.", "Top of receipt"),
            ("Invoice number", "INV-2026-04812", "Document header"),
            ("Invoice date", "2026-04-15", "Document header"),
            ("Due date", "2026-05-15", "Net-30 from invoice date"),
            ("Currency", "USD", "Symbol parsed: $"),
            ("Subtotal", "241.49", "Sum of 4 line items"),
            ("Tax (Sales tax 7%)", "16.90", "Tax row"),
            ("Total", "258.39", "TOTAL DUE row"),
            ("Suggested account", "6000 · Operating Expenses", "AI from entity COA"),
            ("Confidence", "0.95", "AI self-rating"),
        ],
    )
    add_paragraph(doc,
        "After human approval, the engine posts a balanced journal entry. "
        "Thomas Foundation's books are in CHF, so the USD invoice converts "
        "at the FX rate of the invoice date (≈0.9 USD/CHF on 2026-04-15):",
        space_after=8,
    )
    add_example_table(doc,
        ["Account", "Description", "Debit (CHF)", "Credit (CHF)"],
        [
            ("6000 · Operating Expenses", "Office supplies (BlueSpruce)", "232.55", ""),
            ("2000 · Accounts Payable", "BlueSpruce Office Supplies", "", "232.55"),
        ],
        footer=["", "Totals", "232.55", "232.55"],
    )
    add_paragraph(doc,
        "Note that debits equal credits — the foundation of double-entry. "
        "The original USD amount stays linked to the bill record, but the "
        "ledger posts in the entity's functional currency. Both currency "
        "values are visible on the JE detail page.",
        space_after=14,
    )

    # Example 2: Bank line categorization
    add_heading(doc, "5.2 Bank line — wire received from a customer", level=2)
    add_paragraph(doc,
        "A CHF wire of 7,250 lands in the bank account. The bookkeeper "
        "opens the bank-feed page; AI suggests the offset; one click "
        "creates the journal entry.",
        space_after=8,
    )
    add_paragraph(doc, "AI suggestion panel:", bold=True, size=10, space_after=2)
    add_example_table(doc,
        ["Field", "Value"],
        [
            ("Bank line", "+7,250.00 CHF · Sterling Wealth — Q2 advisory fee"),
            ("AI-suggested offset", "4100 · Service Revenue"),
            ("Reasoning", "Positive deposit with vendor name matching open invoice INV-Q2-0042"),
            ("Confidence", "0.91"),
            ("Model", "claude-haiku-4-5"),
        ],
    )
    add_paragraph(doc,
        "The reviewer accepts the suggestion. The engine writes:",
        space_after=8,
    )
    add_example_table(doc,
        ["Account", "Description", "Debit (CHF)", "Credit (CHF)"],
        [
            ("1010 · Operating Bank", "Wire received — Sterling Wealth", "7,250.00", ""),
            ("4100 · Service Revenue", "Q2 advisory fee", "", "7,250.00"),
        ],
        footer=["", "Totals", "7,250.00", "7,250.00"],
    )
    add_paragraph(doc,
        "The journal entry auto-submits for approval, lands on the daily "
        "Approvals queue, and posts when a second human signs off. The "
        "bank line status flips from 'new' to 'matched' once the JE posts.",
        space_after=14,
    )

    # Example 3: Audit trail
    add_heading(doc, "5.3 What the audit trail looks like", level=2)
    add_paragraph(doc,
        "Every action in the system writes an immutable audit event. Below "
        "is the trace for the bill in 5.1, exactly as it appears on the "
        "JE detail page:",
        space_after=8,
    )
    add_example_table(doc,
        ["Time (UTC)", "Actor", "Action", "Detail"],
        [
            ("2026-04-15 14:23:01", "claude-haiku-4-5 (AI)",
             "Drafted journal entry",
             "Confidence 0.95 · model = claude-haiku-4-5"),
            ("2026-04-15 14:23:42", "Demo Admin (human)",
             "Created bill from extraction",
             "Bill BILL-000003 · status = draft"),
            ("2026-04-15 14:24:18", "Demo Admin (human)",
             "Submitted for approval",
             "Status: draft → pending approval"),
            ("2026-04-15 14:31:55", "Sarah Approver (human)",
             "Approved bill",
             "Status: pending approval → approved · accrual JE auto-posts"),
            ("2026-04-15 14:31:55", "(system)",
             "Posted journal entry JE-000004",
             "DR 6000 / CR 2000 · 232.55 CHF · period 2026-04 attached"),
        ],
    )
    add_paragraph(doc,
        "The submitter (Demo Admin) and approver (Sarah Approver) are "
        "different humans — segregation of duties enforced. The AI is "
        "recorded as a distinct actor type, never hidden behind the human "
        "who accepted the suggestion. This entire trail is visible to "
        "auditors with a single page export.",
        space_after=12,
    )

    add_review_box(doc, "any example missing or worth replacing?")

    # ── Section: Architecture principles ──────────────────────────
    add_heading(doc, "6. Architecture principles", level=1)

    principles = [
        ("AI proposes, humans approve.",
         "No record posts to the ledger without a human review action. "
         "AI features can be turned off entirely; the underlying manual "
         "workflows always work."),
        ("Segregation of duties enforced at the engine.",
         "The submitter of a journal entry cannot approve it. Enforcement "
         "is in the service layer — it cannot be disabled by reconfiguring "
         "roles or permissions."),
        ("Every cent traceable to its source.",
         "Reports drill to journal entries, journal entries drill to source "
         "documents, source documents drill to the AI metadata that drafted "
         "them. There is no opaque transformation step."),
        ("Multi-entity, multi-currency, multi-standard.",
         "Each entity has its own functional currency, reporting currency, "
         "and accounting standard (IFRS / US-GAAP / UK-GAAP). FX rates are "
         "stored as a time series; conversions use the most recent rate "
         "on or before the journal entry date."),
        ("Privacy-first by default.",
         "All AI features can run on local infrastructure with nothing "
         "leaving the network. Cloud AI is opt-in per organisation, with "
         "the choice surfaced clearly to the operator."),
        ("Approval state machine is uniform.",
         "Bills, invoices, journal entries, and bank-feed categorisations "
         "all follow the same draft → pending → approved → posted lifecycle. "
         "A single Approvals queue surfaces them together."),
    ]
    for lead, body in principles:
        add_bullet(doc, body, bold_lead=lead)

    add_review_box(doc, "any principle missing or worded wrong?")

    # ── Section: Outstanding work ────────────────────────────────
    add_heading(doc, "7. Where review would help", level=1)
    add_paragraph(doc,
        "The engine is functional end-to-end, but four areas would benefit "
        "from your judgement before they are finalised:",
    )

    open_questions = [
        ("Foreign-exchange rate sourcing.",
         "Today rates are entered manually. The plan is to auto-fetch from "
         "the European Central Bank reference feed daily, with manual override "
         "preserved. We would value your view on whether ECB is the right "
         "default for the family-office context, or whether a different "
         "authoritative source is preferred."),
        ("Period-close workflow.",
         "Closing a period locks new postings into it. Reopening requires "
         "explicit permission. We have not yet defined the exact "
         "month-end / quarter-end / year-end checklist — recurring "
         "recognition entries, FX revaluation, intercompany settlement. "
         "Your standard close checklist would help."),
        ("Account hierarchy in reports.",
         "Reports currently show a flat list of accounts under each section. "
         "Most family-office books expect a hierarchy "
         "(Assets > Current > Cash > Operating Bank). Wiring the hierarchy "
         "into the reports is straightforward; we would like your preferred "
         "depth and the parent-child grouping rules."),
        ("Comparative period column on P&L.",
         "Adding a 'this period vs same period prior year' column on the "
         "income statement is a common accountant ask. Whether it should "
         "be a standing column or an opt-in toggle is a judgement call."),
    ]
    for lead, body in open_questions:
        add_bullet(doc, body, bold_lead=lead)

    add_review_box(doc, "what should be on this list that isn't?")

    # ── Section: FAQ ──────────────────────────────────────────────
    add_heading(doc, "8. Frequently asked questions", level=1)
    add_paragraph(doc,
        "Questions we've heard from advisors, accountants, and family-office "
        "principals during early conversations. Answers reflect what the "
        "engine does today, not promises.",
        space_after=8,
    )

    add_faq(doc,
        "How is Beakon different from Xero, QuickBooks, or Sage?",
        "Those products are SaaS bookkeeping built before useful AI existed. "
        "Beakon is built around the AI loop from the ground up: every intake "
        "(bills, invoices, bank lines) is a draft proposed by AI and approved "
        "by a human, with the model and confidence preserved on the audit "
        "trail. The output reports are the same, but the inputs take seconds "
        "instead of minutes per record. Multi-entity is first-class — five "
        "entities are handled the same way as one — which family offices "
        "typically lose to Xero's per-entity pricing model.",
    )

    add_faq(doc,
        "What happens if the AI gets something wrong?",
        "The reviewer sees the AI's confidence score (under 70% is flagged "
        "yellow; under 40% red), the suggested-account reasoning, and the "
        "accounting-standard explanation on every draft. They can edit any "
        "field before approving. The audit trail records both the AI's "
        "original suggestion AND the human's final decision, so wrong "
        "suggestions never disappear silently — they're visible to anyone "
        "auditing the books later.",
    )

    add_faq(doc,
        "Can I run this without sending data to Anthropic?",
        "Yes. Beakon's AI features have two backends: a local mode (open-source "
        "models running on the office's machine, nothing leaves the network) "
        "and Anthropic Claude (sent to api.anthropic.com over TLS). The choice "
        "is configurable per organisation. Local mode is the default; it "
        "produces good results on clean documents and stays slightly behind "
        "Claude on hard scans or cryptic memos.",
    )

    add_faq(doc,
        "How does multi-currency work?",
        "Each entity has a functional currency (CHF for the foundation, USD "
        "for a US holdco, etc.). Documents in any currency are accepted; "
        "the engine converts at the FX rate of the document's date and "
        "posts in functional. FX rates are stored as a time series — the "
        "engine looks up the most recent rate on or before the journal "
        "entry date. Reports render in either functional or a chosen "
        "reporting currency. Month-end FX revaluation produces a separate "
        "journal entry that hits Realised / Unrealised FX accounts.",
    )

    add_faq(doc,
        "How does this scale to multiple entities?",
        "Every record (account, bill, invoice, journal entry) carries an "
        "entity reference. Reports run in single-entity mode (just one "
        "set of books) or consolidated mode (all entities in one currency, "
        "with intercompany eliminations applied). Each entity carries its "
        "own functional currency, accounting standard, and chart of "
        "accounts; shared accounts (e.g. a holding-company-wide AR control) "
        "are also supported.",
    )

    add_faq(doc,
        "What is the segregation-of-duties enforcement?",
        "The user who submits a journal entry for approval cannot also "
        "approve it. The engine refuses the action at the service layer — "
        "this cannot be bypassed by reconfiguring roles. For solo dev / "
        "platform support, a 'break-glass' override exists for users with "
        "platform-superuser status; every override is tagged in the audit "
        "log so it remains visible. For your end-customers, no override "
        "exists — submitter ≠ approver, always.",
    )

    add_faq(doc,
        "Where are the AI features used today?",
        "Four places: (1) bill OCR + extraction, where receipts become "
        "draft bills with suggested expense accounts; (2) invoice OCR + "
        "extraction, the same flow on the AR side with revenue accounts; "
        "(3) bank-line categorisation, where the offset account on each "
        "bank transaction is suggested; and (4) executive commentary on "
        "the financial reports, written as 3-4 sentences of prose. All "
        "four are opt-in and switchable backends; the manual flow always "
        "works without the AI.",
    )

    add_faq(doc,
        "Can I export everything?",
        "Yes. Every report exports to CSV with one click; the format is "
        "Excel-compatible with proper UTF-8 encoding so umlauts and other "
        "non-ASCII characters render correctly. Journal entries export to "
        "CSV in standard formats accountants expect. The audit log exports "
        "as a date-bracketed CSV. There is no vendor lock-in: if you ever "
        "leave Beakon, you take the data with you.",
    )

    add_faq(doc,
        "What about month-end and year-end close?",
        "Closing a period marks it locked — new journal entries cannot land "
        "in a closed period without a reopen action by an authorised user. "
        "The close workflow runs FX revaluation, recurring recognition "
        "entries (deferrals, accruals), and intercompany eliminations as "
        "configurable steps. The exact close checklist is one of the open "
        "design questions in section 7 — your standard close runbook would "
        "shape this directly.",
    )

    add_faq(doc,
        "What does a typical month cost to run?",
        "The engine itself is software you operate — no per-user, per-entity, "
        "or per-transaction fees from Beakon. The only running cost is the "
        "AI inference: roughly €0.001–0.005 per AI call when using Claude "
        "Haiku, and €0 when using local mode. For a family office processing "
        "200 bills, 100 invoices, and 1,000 bank transactions a month, that "
        "is approximately €1–5/month total in cloud AI costs — orders of "
        "magnitude below what a Xero+Dext+ApprovalMax stack would charge.",
    )

    add_faq(doc,
        "How do I trust the numbers?",
        "Five independent reports (Trial Balance, Balance Sheet, P&L, Cash "
        "Flow, AP/AR Aging) each compute from the same posted journal "
        "entries, and each carries a balance / reconciliation badge. If "
        "the badges are green, the engine is internally consistent. If any "
        "badge turns red, the engine surfaces the exact source of the "
        "disagreement. The trial-balance and balance-sheet equations are "
        "fundamental — a correct accounting system can never silently break "
        "them, and an incorrect one cannot hide the break.",
    )

    add_review_box(doc, "any FAQ that is wrong, or missing, or worded poorly?")

    # ── Section: How to mark this up ─────────────────────────────
    add_heading(doc, "9. How to mark this document", level=1)
    add_paragraph(doc,
        "Each section ends with a cream-coloured Thomas's notes box. Please "
        "use these for written comments — they are intentionally blank and "
        "expandable. You can also use Word's standard Track Changes mode and "
        "Comments — anything that survives a round-trip back to us is fair "
        "game.",
    )
    add_paragraph(doc,
        "Once you return this document, we will lock the version, address "
        "every comment in writing, and re-issue. The result becomes the "
        "canonical Beakon engine specification for the next quarter.",
        space_after=12,
    )

    add_callout(doc,
        "This is a draft. Nothing in it is final. The engine is real and "
        "demonstrable today, but every sentence in this document is "
        "negotiable. Mark the document up and send it back — that is "
        "exactly the workflow Beakon is built around: AI proposes, "
        "humans approve.",
    )

    # ── Save ───────────────────────────────────────────────────────
    doc.save(str(DOCX_PATH))


if __name__ == "__main__":
    build()
    size = DOCX_PATH.stat().st_size
    print(f"Wrote {DOCX_PATH}  ({size:,} bytes)")
